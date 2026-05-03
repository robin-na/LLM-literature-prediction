from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Styles ────────────────────────────────────────────────────────────────────
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(11)

for h_name, sz in [('Heading 1', 13), ('Heading 2', 12), ('Heading 3', 11)]:
    s = doc.styles[h_name]
    s.font.name  = 'Times New Roman'
    s.font.size  = Pt(sz)
    s.font.bold  = True
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.space_before = Pt(12)
    s.paragraph_format.space_after  = Pt(4)

def add_paragraph(text='', style='Normal', bold=False, italic=False, space_before=None, space_after=None):
    p = doc.add_paragraph(style=style)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
    return p

def add_mixed(parts, style='Normal', space_before=None, space_after=None):
    """parts = list of (text, bold, italic, mono)"""
    p = doc.add_paragraph(style=style)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    for text, bold, italic, mono in parts:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        if mono:
            run.font.name = 'Courier New'
            run.font.size = Pt(9.5)
    return p

def add_code_block(lines):
    """Shaded monospace paragraph for CLI commands."""
    text = '\n'.join(lines)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    # light gray shading on paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F2F2F2')
    pPr.append(shd)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent   = Cm(0.75 + level * 0.5)
    p.paragraph_format.space_before  = Pt(2)
    p.paragraph_format.space_after   = Pt(2)
    p.add_run(text)
    return p

def add_numbered(text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.add_run(text)
    return p

def set_col_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'),    str(int(width_cm * 567)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def shade_row(row, fill='E8E8E8'):
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  fill)
        tcPr.append(shd)

def cell_text(cell, text, bold=False, mono=False, sz=10):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    if mono:
        run.font.name = 'Courier New'
    run.font.size = Pt(sz)

# ═════════════════════════════════════════════════════════════════════════════
# TITLE
# ═════════════════════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_after = Pt(16)
r = title_p.add_run('Supplementary Information: Automated Extraction of Experimental Parameters from the Public Goods Game Literature')
r.bold      = True
r.font.name = 'Times New Roman'
r.font.size = Pt(14)

# ═════════════════════════════════════════════════════════════════════════════
# S1 OVERVIEW
# ═════════════════════════════════════════════════════════════════════════════
doc.add_heading('S1. Overview', level=1)

add_paragraph(
    'We developed a fully automated pipeline to extract structured experimental parameters '
    'from empirical papers on Public Goods Games (PGGs). The goal was to construct a '
    'machine-readable database of experimental conditions — one record per treatment arm — '
    'covering experimental design choices, independent and dependent variables, participant '
    'characteristics, and data provenance. The resulting dataset enables systematic '
    'comparisons across the PGG literature and serves as the empirical foundation for the '
    'prediction task described in the main paper.'
)
add_paragraph(
    'The pipeline uses a large language model (LLM) to read each paper\'s full text in '
    'Markdown format and populate a predefined schema. Extraction is performed at the '
    'condition level: if a paper reports four treatment arms (e.g., a 2×2 design), the '
    'pipeline produces four rows, one per arm. Field definitions, encoding rules, and worked '
    'examples are supplied in the prompt; the model is not asked to infer, impute, or '
    'extrapolate beyond what the paper explicitly states.'
)
add_mixed([
    ('All code is available in the associated GitHub repository (', False, False, False),
    ('https://github.com/robin-na/LLM-literature-prediction', False, False, False),
    ('), under ', False, False, False),
    ('batch_processing/', False, False, True),
    ('.', False, False, False),
])

# ═════════════════════════════════════════════════════════════════════════════
# S2 CORPUS
# ═════════════════════════════════════════════════════════════════════════════
doc.add_heading('S2. Corpus', level=1)
doc.add_heading('Paper Selection', level=2)
add_paragraph(
    'The corpus consists of 810 papers identified through Web of Science searches for '
    'empirical and experimental PGG studies. Papers were selected based on title and abstract '
    'screening for relevance to the public goods contribution game paradigm, including variants '
    'with punishment, reward, communication, and other institutional mechanisms.'
)

doc.add_heading('Preparation', level=2)
add_mixed([
    ('Each paper was converted from PDF to Markdown (', False, False, False),
    ('.md', False, False, True),
    (') format and stored in a shared Google Drive folder (', False, False, False),
    ('papers_markdown', False, False, True),
    ('). The conversion preserves text, tables, and equations in a format the LLM can read '
     'directly. The script ', False, False, False),
    ('download_papers_md.py', False, False, True),
    (' authenticates via OAuth 2.0 and downloads each file by matching the DOI-based filename. '
     'Files already present locally are skipped, making the download step idempotent.',
     False, False, False),
])

doc.add_heading('Scale', level=2)
add_bullet('Papers submitted for extraction: 810')
add_bullet('Papers with at least one extractable experimental condition: 808')
add_bullet('Papers excluded (returned empty conditions): 2 — both were literature reviews or meta-analyses with no original experimental data')
add_bullet('Total condition rows extracted: 3,910')

# ═════════════════════════════════════════════════════════════════════════════
# S3 PIPELINE
# ═════════════════════════════════════════════════════════════════════════════
doc.add_heading('S3. Extraction Pipeline', level=1)
doc.add_heading('Architecture', level=2)

add_paragraph('Figure S1 illustrates the end-to-end extraction workflow.', space_after=6)

# ── Pipeline flowchart table ──────────────────────────────────────────────────
flow_steps = [
    'PDF papers on Google Drive',
    'download_papers_md.py  |  OAuth2 fetch → .md files',
    'build_batch_input.py  |  Schema + instructions + full text → JSONL',
    'batch-submit  |  Upload JSONL → OpenAI Batch API',
    'OpenAI Batch API  |  GPT-4.1, T=0, JSON output  |  (~50% cheaper, up to 24 h)',
    'batch-collect  |  Download results → parse experiments',
    'fill_paper_level_counts()  |  Compute number_IVs, number_DVs',
    'Excel workbook  |  extractions sheet (3,910 rows × 136 columns)',
]

for i, step in enumerate(flow_steps):
    parts = step.split('  |  ')
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(10 * 567)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'DCE6F1')
    tcPr.append(shd)
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(parts[0])
    run.bold      = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    for sub in parts[1:]:
        p.add_run('\n')
        r2 = p.add_run(sub)
        r2.font.name = 'Courier New'
        r2.font.size = Pt(9)

    if i < len(flow_steps) - 1:
        arrow = doc.add_paragraph('▼')
        arrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
        arrow.paragraph_format.space_before = Pt(0)
        arrow.paragraph_format.space_after  = Pt(0)
        arrow.runs[0].font.size = Pt(10)

caption = add_paragraph(
    'Figure S1. End-to-end extraction pipeline. Each box corresponds to a discrete processing '
    'step. The OpenAI Batch API processes all 810 papers asynchronously; results are collected '
    'once all jobs are complete.',
    space_before=6
)
caption.runs[0].italic = True

doc.add_heading('Model and API', level=2)
add_mixed([
    ('Extraction uses ', False, False, False),
    ('GPT-4.1', True, False, False),
    (' (OpenAI) at temperature 0 to ensure deterministic output. All requests are submitted via the ', False, False, False),
    ('OpenAI Batch API', True, False, False),
    (' (', False, False, False),
    ('/v1/responses', False, False, True),
    (' endpoint with ', False, False, False),
    ('response_format: {type: json_object}', False, False, True),
    ('), which processes jobs asynchronously at approximately 50% of the real-time API cost and returns results within 24 hours.',
     False, False, False),
])
add_paragraph('Each request consists of:')
add_numbered('A system prompt specifying the extraction role and JSON output constraint.')
add_numbered('A user prompt comprising (a) the output schema definition, (b) field-by-field instructions and worked examples, and (c) the full paper text.')

doc.add_heading('Prompt Design', level=2)
add_paragraph('Non-inference rule.  The instruction set opens with a strict non-inference rule:', space_after=2)

quote = doc.add_paragraph(
    '"Only report a value when the paper explicitly states it or when it is directly and '
    'unambiguously computable from reported numbers. Do NOT infer values from silence, '
    'convention, or reasonable assumption."'
)
quote.paragraph_format.left_indent  = Cm(1.2)
quote.paragraph_format.right_indent = Cm(1.2)
quote.paragraph_format.space_before = Pt(2)
quote.paragraph_format.space_after  = Pt(6)
quote.runs[0].italic = True

add_paragraph(
    'This prevents the model from filling in plausible but unverified values based on domain conventions.',
    space_after=6
)

add_paragraph('N/A versus N/R.  Two sentinel values distinguish different types of missingness:', space_after=2)
add_mixed([
    ('N/A', True, False, False),
    (' — the field is not applicable to this condition (e.g., ', False, False, False),
    ('CONFIG_punishmentCost', False, False, True),
    (' when no punishment mechanism exists).', False, False, False),
], style='List Bullet')
add_mixed([
    ('N/R', True, False, False),
    (' — the field is applicable but the paper does not report the value.', False, False, False),
], style='List Bullet')
add_paragraph('Coding rules specify which sentinel to use for each field, preventing false zeros.', space_before=4)

add_paragraph(
    'Condition-level instruction.  The model is instructed to produce one JSON object per '
    'experimental condition within the experiments array. Papers with multi-arm designs produce '
    'multiple objects; single-condition papers produce one.'
)
add_paragraph(
    'Reason and confidence.  For each field X, the model also returns X_reason (a brief '
    'justification quoting or paraphrasing the paper) and X_confidence (a self-assessed numeric '
    'confidence on 0–1). These audit columns facilitate quality review.'
)

# ═════════════════════════════════════════════════════════════════════════════
# S4 SCHEMA
# ═════════════════════════════════════════════════════════════════════════════
doc.add_heading('S4. Extraction Schema', level=1)
add_paragraph(
    'The schema comprises 45 fields organised into five groups. The complete output has '
    '136 columns per row: custom_id plus three columns per field (value, reason, confidence). '
    'Table S1 lists all fields with their type, granularity, and definition.'
)
add_mixed([
    ('Granularity notation.  ', True, False, False),
    ('Condition', False, True, False),
    (' fields may differ across treatment arms within a paper. ', False, False, False),
    ('Paper', False, True, False),
    (' fields take the same value for every row belonging to the same paper; they are '
     'computed or extracted once and propagated. Fields marked † are computed programmatically '
     'after LLM extraction (not asked to the model).', False, False, False),
])

add_paragraph('Table S1. Complete extraction schema.', bold=True, space_before=10, space_after=4)

# Schema table data: (field, type, granularity, definition)
schema_rows = [
    # section headers use ('HEADER', '', '', '') sentinel
    ('CONDITION IDENTIFICATION', None, None, None),
    ('data_id', 'string', 'condition', 'Label the paper uses for this condition (e.g., "Experiment 1 – Punishment Treatment")'),
    ('indep_var', 'string', 'condition', 'Independent variable(s) and their value(s) for this row'),
    ('METHOD CLASSIFICATION', None, None, None),
    ('METHOD_empirical', 'bool', 'condition', 'True if the study uses human participants'),
    ('METHOD_experiment', 'bool', 'condition', 'True if controlled experiment; false if observational'),
    ('METHOD_lab', 'bool', 'condition', 'True if lab-based; false if field experiment'),
    ('METHOD_simulation', 'bool', 'condition', 'True if agent-based / computer simulation with no human subjects'),
    ('METHOD_analytical', 'bool', 'condition', 'True if formal mathematical model (no human subjects)'),
    ('EXPERIMENTAL CONFIGURATION', None, None, None),
    ('CONFIG_playerCount', 'int', 'condition', 'Number of strategic players per group. Teams count as one player; session headcount is not reported.'),
    ('CONFIG_numRounds', 'int', 'condition', 'Number of game rounds or periods'),
    ('CONFIG_allOrNothing', '0/1', 'condition', '1 = continuous contribution choice (any amount 0–endowment); 0 = binary all-or-nothing. N/A if not described. Note: 1 denotes the more flexible, continuous case.'),
    ('CONFIG_defaultContribProp', '0–1', 'condition', 'Fraction of endowment placed in the public fund by default. 0 = standard VCM (private account default); 1 = opt-out game. N/A if not stated.'),
    ('CONFIG_MPCR', 'float', 'condition', 'Marginal per-capita return: group multiplier divided by group size'),
    ('CONFIG_chat', '0/1', 'condition', '1 = unrestricted free-form communication explicitly permitted; 0 = explicitly prohibited; N/A if not mentioned'),
    ('CONFIG_showOtherSummaries', '0/1', 'condition', '1 = participants receive post-round feedback on others\' contributions or earnings; N/A if not mentioned'),
    ('CONFIG_showPunishmentId', '0/1', 'condition', '1 = punished player can identify the punisher; 0 = anonymous punishment; N/A if no punishment or not stated'),
    ('CONFIG_showRewardId', '0/1', 'condition', '1 = reward recipient can identify the rewarder; N/A if no reward or not stated'),
    ('CONFIG_showNRounds', '0/1', 'condition', '1 only if paper explicitly states the round count is displayed to participants. N/A if not stated — never coded 0 from silence.'),
    ('CONFIG_punishmentExists', '0/1', 'condition', '1 only if a punishment mechanism is explicitly described for this condition; N/A if not mentioned'),
    ('CONFIG_punishmentCost', 'float', 'condition', 'Tokens spent by the punisher per unit of punishment assigned. N/A if no punishment.'),
    ('CONFIG_punishmentTech', 'float', 'condition', 'Reduction in target\'s payoff per unit of punishment received (the "effectiveness" ratio). N/A if no punishment.'),
    ('CONFIG_rewardExists', '0/1', 'condition', '1 only if a reward mechanism is explicitly described; N/A if not mentioned'),
    ('CONFIG_rewardCost', 'float', 'condition', 'Tokens spent by the rewarder per unit of reward. N/A if no reward.'),
    ('CONFIG_rewardTech', 'float', 'condition', 'Increase in recipient\'s payoff per unit of reward received. N/A if no reward.'),
    ('CONFIG_endowment', 'float', 'condition', 'Initial token allocation per player per round'),
    ('INDEPENDENT AND DEPENDENT VARIABLES', None, None, None),
    ('IVs', 'JSON array', 'condition', 'Snake_case names of experimental factors varied across conditions in this paper. Fixed parameters (e.g., endowment when constant across arms) are excluded. Example: ["punishment_mechanism", "group_size"].'),
    ('number_IVs †', 'int', 'paper', 'Count of unique IV names across all conditions in the paper (union across conditions). Same value in every row for a given paper.'),
    ('DVs', 'JSON array', 'condition', 'Snake_case names of outcome measures reported for this condition. Example: ["individual_contribution", "efficiency"].'),
    ('DVs_Definitions', 'JSON object', 'condition', 'Maps each DV name to its paper-specific definition. Keys must match the DVs array exactly.'),
    ('number_DVs †', 'int', 'paper', 'Count of unique DV names across all conditions in the paper (union across conditions). Same value in every row.'),
    ('DV_efficiencyReported', '0/1', 'paper', '1 if efficiency (actual group payoff / maximum cooperative payoff) is reported anywhere in the paper; 0 only if it never appears. Same value in every row.'),
    ('STUDY CONTEXT', None, None, None),
    ('source_data', 'string', 'paper', '"Internal": data collected by the authors for this study (lab, online, or field). "External": data borrowed from another experiment or archival dataset (reanalysis). Same value in every row.'),
    ('participant_country', 'string', 'condition', 'Country where the experiment was conducted. N/R if not stated.'),
    ('participant_age', 'string', 'condition', 'Age descriptor (e.g., "student", "18–35 years"). N/R if not stated.'),
    ('participant_gender', 'string', 'condition', 'Gender composition: "mixed", "male only", "female only", or N/R.'),
    ('participant_education', 'string', 'condition', 'Education level (e.g., "undergraduate students", "general population"). N/R if not stated.'),
    ('experiment_environment', 'categ.', 'condition', 'One of: Online, On site, Field experiment, Observational, No human'),
    ('other_game_info', 'string', 'condition', 'Free-text field for additional design details not captured by other fields'),
]

tbl = doc.add_table(rows=1, cols=4)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

# Column widths in cm
col_widths = [4.0, 1.6, 2.2, 7.8]

# Header row
hdr = tbl.rows[0]
shade_row(hdr, 'C0C0C0')
for i, (txt, w) in enumerate(zip(['Field', 'Type', 'Granularity', 'Definition'], col_widths)):
    set_col_width(hdr.cells[i], w)
    cell_text(hdr.cells[i], txt, bold=True, sz=10)

for field, ftype, gran, defn in schema_rows:
    if ftype is None:
        # Section header row
        row = tbl.add_row()
        shade_row(row, 'BDD7EE')
        merged = row.cells[0].merge(row.cells[1]).merge(row.cells[2]).merge(row.cells[3])
        cell_text(merged, field, bold=True, sz=10)
    else:
        row = tbl.add_row()
        for i, w in enumerate(col_widths):
            set_col_width(row.cells[i], w)
        cell_text(row.cells[0], field, mono=True, sz=9)
        cell_text(row.cells[1], ftype, sz=10)
        cell_text(row.cells[2], gran, sz=10)
        cell_text(row.cells[3], defn, sz=10)

cap = add_paragraph(
    'Table S1. Complete extraction schema. Fields marked † are computed programmatically '
    'after LLM extraction (not asked to the model).',
    space_before=4
)
cap.runs[0].italic = True

# ═════════════════════════════════════════════════════════════════════════════
# S5 POST-PROCESSING
# ═════════════════════════════════════════════════════════════════════════════
doc.add_heading('S5. Post-Processing', level=1)
add_mixed([
    ('Two fields — ', False, False, False),
    ('number_IVs', False, False, True),
    (' and ', False, False, False),
    ('number_DVs', False, False, True),
    (' — are computed programmatically after LLM extraction rather than being asked of the '
     'model directly. This design choice avoids a systematic LLM error in which models count ',
     False, False, False),
    ('levels', False, True, False),
    (' within a variable rather than unique variable ', False, False, False),
    ('names', False, True, False),
    (' (e.g., treating the binary variable "punishment: yes/no" as two independent variables '
     'rather than one).', False, False, False),
])
add_mixed([
    ('The computation is performed by ', False, False, False),
    ('fill_paper_level_counts()', False, False, True),
    (' in ', False, False, False),
    ('extraction_pipeline.py', False, False, True),
    (':', False, False, False),
])

add_numbered('All condition rows for a paper are grouped by paper identifier (custom_id).')
add_numbered('For each paper, the IVs and DVs JSON arrays are parsed from every condition row.')
add_numbered('The union of all IV names and the union of all DV names are accumulated: number_IVs = |∪ IVs across conditions|, number_DVs = |∪ DVs across conditions|.')
add_numbered('The computed integer is written to every row of that paper (paper-level field).')

add_paragraph(
    'The union-based approach correctly handles designs in which some DVs appear only in '
    'specific conditions (e.g., punishment-related outcomes appear only in punishment arms): '
    'such DVs are included in the paper-level count because they appear in at least one condition.',
    space_before=6
)

# ═════════════════════════════════════════════════════════════════════════════
# S6 QUALITY AND SCALE
# ═════════════════════════════════════════════════════════════════════════════
doc.add_heading('S6. Quality and Scale', level=1)
doc.add_heading('Extraction Results', level=2)

add_paragraph('Table S2. Summary statistics for the extracted dataset.', bold=True, space_after=4)

stats = [
    ('Papers submitted', '810'),
    ('Papers with extractable conditions', '808'),
    ('Papers returning empty conditions', '2'),
    ('Total condition rows', '3,910'),
    ('Fields per condition', '45'),
    ('Output columns per row', '136'),
    ('source_data = Internal', '3,862 rows (98.8%)'),
    ('source_data = External', '48 rows (1.2%)'),
]
tbl2 = doc.add_table(rows=1, cols=2)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
shade_row(tbl2.rows[0], 'C0C0C0')
cell_text(tbl2.rows[0].cells[0], 'Metric', bold=True, sz=10)
cell_text(tbl2.rows[0].cells[1], 'Value',  bold=True, sz=10)
set_col_width(tbl2.rows[0].cells[0], 9.0)
set_col_width(tbl2.rows[0].cells[1], 3.0)
for metric, val in stats:
    row = tbl2.add_row()
    set_col_width(row.cells[0], 9.0)
    set_col_width(row.cells[1], 3.0)
    cell_text(row.cells[0], metric, sz=10)
    cell_text(row.cells[1], val,    sz=10)

cap2 = add_paragraph('Table S2. Summary statistics for the extracted dataset.', space_before=4)
cap2.runs[0].italic = True

doc.add_heading('Papers Excluded', level=2)
add_mixed([
    ('Two papers returned ', False, False, False),
    ('{"experiments": []}', False, False, True),
    ('. Manual inspection confirmed both are literature reviews or meta-analyses that '
     'synthesise prior results without reporting original experimental data; they contain '
     'no conditions to extract. No extraction errors occurred.', False, False, False),
])

doc.add_heading('Output File', level=2)
add_mixed([
    ('The extracted dataset is stored in ', False, False, False),
    ('batch_processing/output_xlsx/simple_batch_810papers.xlsx', False, False, True),
    (', sheet ', False, False, False),
    ('extractions', False, False, True),
    ('. Each row corresponds to one experimental condition; each field has three associated '
     'columns (value, reason, confidence score). Arrays (', False, False, False),
    ('IVs', False, False, True),
    (', ', False, False, False),
    ('DVs', False, False, True),
    (') and dictionaries (', False, False, False),
    ('DVs_Definitions', False, False, True),
    (') are stored as JSON strings.', False, False, False),
])

# ═════════════════════════════════════════════════════════════════════════════
# S7 REPRODUCIBILITY
# ═════════════════════════════════════════════════════════════════════════════
doc.add_heading('S7. Reproducibility', level=1)
add_paragraph('The extraction can be reproduced by following these steps:')

add_mixed([
    ('Step 1 — Obtain papers.  ', True, False, False),
    ('Download Markdown files from Google Drive using ', False, False, False),
    ('download_papers_md.py', False, False, True),
    (' with the PDF path list ', False, False, False),
    ('batch_processing/inputs/pdf_paths_800.txt', False, False, True),
    ('.', False, False, False),
])

add_mixed([
    ('Step 2 — Submit batch job.', True, False, False),
])
add_code_block([
    'python batch_processing/extract_papers.py batch-submit \\',
    '  --paper-dir PGG_papers/papers \\',
    '  --paper-ids $(cat batch_processing/inputs/paper_ids_800.txt | tr \'\\n\' \' \') \\',
    '  --model gpt-4.1 \\',
    '  --save-jsonl batch_processing/inputs/batch_input_810papers.jsonl',
])

add_mixed([
    ('Step 3 — Collect results', True, False, False),
    (' (once the batch status reports ', False, False, False),
    ('completed', False, False, True),
    (').', False, False, False),
])
add_code_block([
    'python batch_processing/extract_papers.py batch-collect <batch_id> \\',
    '  --output-xlsx batch_processing/output_xlsx/simple_batch_810papers.xlsx \\',
    '  --save-jsonl batch_processing/inputs/batch_output_810papers.jsonl',
])

add_mixed([
    ('The raw batch output JSONL (', False, False, False),
    ('batch_output_810papers.jsonl', False, False, True),
    (') is retained alongside the Excel file for full traceability.', False, False, False),
])

# ═════════════════════════════════════════════════════════════════════════════
# SAVE
# ═════════════════════════════════════════════════════════════════════════════
out_path = 'docs/formal/SI_extraction.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
