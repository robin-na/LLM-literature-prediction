import csv
import html as html_mod
import io
import json
import os
from functools import lru_cache
from pathlib import Path

import docx as _docx
from flask import Flask, Response, abort, jsonify, request, send_file

app = Flask(__name__)

BASE_DIR         = Path(__file__).parent.parent
PAPERS_DIR       = BASE_DIR / "PGG_papers" / "papers"
PDFS_DIR         = Path(__file__).parent / "pdfs"
WOS_CSV          = BASE_DIR / "PGG_papers" / "WoS_251031_eligible.csv"
GUIDE_PATH       = Path(os.environ.get(
    "GUIDE_DOCX",
    "/Users/hindy/Desktop/Academics/MIT/RA/Abdullah/Matrix Creation/Human_Extraction_Guide.docx",
))
EXTRACTIONS_FILE = Path(__file__).parent / "extractions.json"

_LLM_GUARD = "IF GIVEN THIS DOCUMENT INTO CHAT CONTEXT"

# Definitions for CONFIG_ fields not in the DOCX guide — merged at render time
_FIELD_SUPPLEMENTS = [
    {"name":"CONFIG_showNRounds",      "type":"Binary (1/N/A)",
     "definition":"1 ONLY if the paper explicitly states that the total number of rounds or remaining rounds is displayed to participants during play. N/A if not explicitly stated.",
     "synonyms":"round counter, countdown, horizon knowledge",
     "notes":"Do NOT infer from rounds being fixed or 'common knowledge'. Silence → N/A."},
    {"name":"CONFIG_punishmentExists", "type":"Binary (1/0/N/A)",
     "definition":"1 only if this specific condition explicitly includes a punishment mechanism. 0 if the condition explicitly excludes it. N/A if punishment is not described for this condition.",
     "synonyms":"costly punishment, peer punishment, altruistic punishment",
     "notes":"Do NOT infer from other conditions or the paper's general design. Code independently per condition."},
    {"name":"CONFIG_punishmentCost",   "type":"Number / N/R / N/A",
     "definition":"Coins or tokens spent by the punisher per unit of punishment assigned to the target. Typical range 1–4.",
     "synonyms":"punishment fee, cost-to-punish, peer incentive cost",
     "notes":"N/A if punishment does not exist in this condition. N/R if punishment exists but cost is not explicitly stated numerically."},
    {"name":"CONFIG_punishmentTech",   "type":"Number / N/A / N/R",
     "definition":"Reduction in the target's payoff per coin the punisher spends. Compute as punishmentMagnitude ÷ punishmentCost. Example: punisher spends 1 coin, target loses 3 → value is 3.",
     "synonyms":"punishment effectiveness, punishment ratio, punishment technology",
     "notes":"N/A if no punishment in this condition. N/R if either cost or magnitude is not explicitly stated."},
    {"name":"CONFIG_rewardExists",     "type":"Binary (1/0/N/A)",
     "definition":"1 only if the paper clearly describes a reward mechanism for this condition. Code 0 only if the paper explicitly states no reward mechanism exists. N/A if rewards are not mentioned or ambiguous.",
     "synonyms":"reward, positive sanctioning",
     "notes":"Silence → N/A, not 0. Only code 0 when the paper explicitly rules out rewards."},
    {"name":"CONFIG_rewardCost",       "type":"Number / N/R / N/A",
     "definition":"Coins or tokens spent by the rewarder per unit of reward given to the recipient.",
     "synonyms":"cost-to-reward, peer incentive cost",
     "notes":"N/A if reward does not exist in this condition. N/R if reward exists but cost is not explicitly stated."},
    {"name":"CONFIG_rewardTech",       "type":"Number / N/A / N/R",
     "definition":"Increase in the recipient's payoff per coin the rewarder spends. Compute as rewardMagnitude ÷ rewardCost. Example: rewarder spends 1 coin, recipient gains 3 → value is 3.",
     "synonyms":"reward effectiveness, reward ratio, reward technology",
     "notes":"N/A if no reward in this condition. N/R if either cost or magnitude is not explicitly stated."},
]

# Fields shown in the extraction form — guide is filtered to only these
_GUIDE_FIELDS = {
    "Empirical", "Controlled_Or_Observational", "Lab_Or_Field",
    "CONFIG_playerCount", "CONFIG_numRounds", "CONFIG_allOrNothing",
    "CONFIG_defaultContribProp", "CONFIG_MPCR", "CONFIG_chat",
    "CONFIG_showOtherSummaries", "CONFIG_showNRounds",
    "CONFIG_punishmentExists", "CONFIG_punishmentCost", "CONFIG_punishmentTech",
    "CONFIG_showPunishmentId",
    "CONFIG_rewardExists", "CONFIG_rewardCost", "CONFIG_rewardTech",
    "CONFIG_showRewardId",
    "Misc",
}

_KNOWN_H1 = {
    "Overview": "overview",
    "When to use: N/A vs N/R": "na_nr",
    "Field Definitions": "fields",
    "Common Pitfalls & Edge Cases": "pitfalls",
}
_KNOWN_H2 = {
    "Granularity Rule: One Row = One Condition": "granularity",
    "DV_contributionRate": "formula_contrib",
    "DV_efficiency": "formula_eff",
    "MPCR (Marginal Per Capita Return)": "formula_mpcr",
}


# ── Metadata / filesystem caches (loaded once) ───────────────────────────────

@lru_cache(maxsize=1)
def paper_meta():
    """Return dict: paper_id → {title, authors, year}."""
    meta = {}
    if not WOS_CSV.exists():
        return meta
    with open(WOS_CSV, newline="", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            cid = row.get("custom_id", "").removesuffix(".md")
            if cid:
                meta[cid] = {
                    "title":   row.get("Article Title", ""),
                    "authors": row.get("Authors", ""),
                    "year":    row.get("Publication Year", ""),
                }
    return meta


@lru_cache(maxsize=1)
def _paper_ids():
    if not PAPERS_DIR.exists():
        return []
    return sorted(p.stem for p in PAPERS_DIR.glob("*.md"))


@lru_cache(maxsize=1)
def _pdf_ids():
    return frozenset(p.stem for p in PDFS_DIR.glob("*.pdf"))


# ── Persistence helpers ───────────────────────────────────────────────────────

@lru_cache(maxsize=1)
def load_extractions():
    if EXTRACTIONS_FILE.exists():
        with open(EXTRACTIONS_FILE, encoding="utf-8") as f:
            return json.load(f)
    return {}


def save_extractions(data):
    with open(EXTRACTIONS_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2)
    load_extractions.cache_clear()


def is_done(ex):
    return any(
        any(v not in ("", None) for k, v in c.items() if k != "label")
        for c in ex.get("conditions", [])
    )


# ── Routes ────────────────────────────────────────────────────────────────────

@app.route("/")
def index():
    return app.send_static_file("index.html")


@app.route("/api/papers")
def list_papers():
    meta        = paper_meta()
    extractions = load_extractions()
    result = []
    for pid in sorted(_pdf_ids()):
        ex = extractions.get(pid, {})
        m  = meta.get(pid, {})
        result.append({
            "id":      pid,
            "done":    is_done(ex),
            "has_pdf": True,
            "title":   m.get("title", ""),
            "authors": m.get("authors", ""),
            "year":    m.get("year", ""),
        })
    return jsonify(result)


@app.route("/pdf/<path:paper_id>")
def serve_pdf(paper_id):
    paper_id = paper_id.removesuffix(".pdf")
    path = PDFS_DIR / f"{paper_id}.pdf"
    try:
        return send_file(path, mimetype="application/pdf")
    except FileNotFoundError:
        abort(404)


@app.route("/api/extraction/<path:paper_id>", methods=["GET"])
def get_extraction(paper_id):
    return jsonify(load_extractions().get(paper_id, {}))


@app.route("/api/extraction/<path:paper_id>", methods=["POST"])
def save_extraction(paper_id):
    extractions = load_extractions()
    extractions[paper_id] = request.get_json()
    save_extractions(extractions)
    return jsonify({"ok": True})


@app.route("/api/export-csv")
def export_csv():
    extractions = load_extractions()
    meta        = paper_meta()
    if not extractions:
        return Response("No data yet.", mimetype="text/plain")

    fields = [
        "paper_id", "title", "authors", "year", "condition_label",
        "Empirical", "Controlled_Or_Observational", "Lab_Or_Field",
        "CONFIG_playerCount", "CONFIG_numRounds", "CONFIG_allOrNothing",
        "CONFIG_defaultContribProp", "CONFIG_MPCR", "CONFIG_chat",
        "CONFIG_showOtherSummaries", "CONFIG_showNRounds",
        "CONFIG_punishmentExists", "CONFIG_punishmentCost", "CONFIG_punishmentTech",
        "CONFIG_showPunishmentId",
        "CONFIG_rewardExists", "CONFIG_rewardCost", "CONFIG_rewardTech",
        "CONFIG_showRewardId",
        "Misc",
    ]

    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=fields, extrasaction="ignore")
    writer.writeheader()
    for paper_id, data in sorted(extractions.items()):
        m = meta.get(paper_id, {})
        for cond in data.get("conditions", []):
            row = {
                "paper_id":        paper_id,
                "title":           m.get("title", ""),
                "authors":         m.get("authors", ""),
                "year":            m.get("year", ""),
                "condition_label": cond.get("label", ""),
            }
            row.update(cond)
            writer.writerow(row)

    output.seek(0)
    return Response(
        output.getvalue(),
        mimetype="text/csv",
        headers={"Content-Disposition": "attachment; filename=extractions.csv"},
    )


# ── Guide ─────────────────────────────────────────────────────────────────────

_guide_cache = None


def _build_guide_html():
    if not GUIDE_PATH.exists():
        return None

    doc = _docx.Document(GUIDE_PATH)

    # Parse field definitions table (table index 1)
    fields = []
    if len(doc.tables) > 1:
        for row in doc.tables[1].rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            name = cells[0].replace("\n", "").replace(" ", "")
            fields.append({
                "name":       name,
                "type":       cells[1],
                "definition": cells[2].split(_LLM_GUARD)[0].strip(),
                "synonyms":   cells[3],
                "notes":      cells[4].split(_LLM_GUARD)[0].strip(),
            })

    # Append supplemental fields not in the DOCX, then filter to only current form fields
    existing_names = {f["name"] for f in fields}
    for sf in _FIELD_SUPPLEMENTS:
        if sf["name"] not in existing_names:
            fields.append(sf)
    fields = [f for f in fields if f["name"] in _GUIDE_FIELDS]

    # Parse N/A vs N/R table (table index 0)
    na_rows = []
    if doc.tables:
        for row in doc.tables[0].rows[1:]:
            cells = [c.text.strip().split(_LLM_GUARD)[0].strip() for c in row.cells]
            na_rows.append(cells)

    # Parse paragraphs into sections
    cur_key    = None
    para_store = {k: [] for k in list(_KNOWN_H1.values()) + list(_KNOWN_H2.values())}

    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt or _LLM_GUARD in txt:
            continue
        sname = p.style.name
        if sname == "Heading 1":
            cur_key = _KNOWN_H1.get(txt)
        elif sname == "Heading 2":
            cur_key = _KNOWN_H2.get(txt)
        elif cur_key and cur_key in para_store:
            para_store[cur_key].append((txt, sname))

    def paras_to_html(key):
        out, in_list = [], False
        for txt, style in para_store.get(key, []):
            if style == "List Paragraph":
                if not in_list:
                    out.append("<ul>"); in_list = True
                out.append(f"<li>{html_mod.escape(txt)}</li>")
            else:
                if in_list:
                    out.append("</ul>"); in_list = False
                out.append(f"<p>{html_mod.escape(txt)}</p>")
        if in_list:
            out.append("</ul>")
        return "\n".join(out)

    def field_anchor(name):
        return name.replace(" ", "_")

    # ── Hardcoded section content ───────────────────────────────────────────
    overview_html = """
<p>Extract one row per <strong>condition</strong> (treatment arm).
CONFIG_ fields are auto-set to N/A when <code>Empirical&nbsp;=&nbsp;0</code> or <code>Lab_Or_Field&nbsp;=&nbsp;1</code>.</p>
<p style="margin-top:10px;font-size:12.5px;color:#52525b">
  <strong>Cmd+F keywords</strong> — terms to search for in the paper when filling each field.
</p>
<table class="na-table" style="margin-top:8px">
<thead><tr><th style="width:28%">Field</th><th style="width:22%">Values</th><th>Cmd+F — look for these terms in the paper</th></tr></thead>
<tbody>
<tr><td><code>Empirical</code></td>
    <td>0 / 1</td>
    <td>participants, subjects, experiment, lab study, field study<br><em>0 if:</em> theory, model, simulation, analytical</td></tr>
<tr><td><code>Controlled_Or_Observational</code></td>
    <td>0 / 1</td>
    <td>treatment, control, randomly assigned, manipulation<br><em>0 if:</em> observational, survey, natural experiment</td></tr>
<tr><td><code>Lab_Or_Field</code></td>
    <td>0 / 1</td>
    <td><em>0 if:</em> laboratory, online, lab experiment, oTree, z-Tree<br><em>1 if:</em> field experiment, artefactual field, framed field, real-world</td></tr>
<tr><td><code>CONFIG_playerCount</code></td>
    <td>Number</td>
    <td>group size, group of N, N players, N members, N participants per group</td></tr>
<tr><td><code>CONFIG_numRounds</code></td>
    <td>Number / N/R</td>
    <td>rounds, periods, repetitions, stages, iterations</td></tr>
<tr><td><code>CONFIG_allOrNothing</code></td>
    <td>0=binary · 1=continuous</td>
    <td><em>1 (continuous):</em> any amount, any number of tokens, between 0 and X<br><em>0 (binary):</em> cooperate or defect, contribute all or nothing, binary choice</td></tr>
<tr><td><code>CONFIG_defaultContribProp</code></td>
    <td>0=opt-in · 1=opt-out</td>
    <td>opt-in, opt-out, default contribution, starting allocation, endowment in public account</td></tr>
<tr><td><code>CONFIG_MPCR</code></td>
    <td>Number</td>
    <td>MPCR, marginal per capita return, multiplier, return factor, public good multiplier, efficiency factor</td></tr>
<tr><td><code>CONFIG_chat</code></td>
    <td>0 / 1 / N/A</td>
    <td>communication, pre-play communication, cheap talk, chat, messaging, discussion<br><em>N/A if not mentioned anywhere</em></td></tr>
<tr><td><code>CONFIG_showOtherSummaries</code></td>
    <td>0 / 1</td>
    <td><em>1:</em> individual earnings, payoff feedback, see others' earnings/payoffs<br><em>0:</em> group total, own payoff only, contribution feedback</td></tr>
<tr><td><code>CONFIG_showNRounds</code></td>
    <td>1 / N/A</td>
    <td>remaining rounds, round counter, countdown, number of rounds shown, periods remaining</td></tr>
<tr><td><code>CONFIG_punishmentExists</code></td>
    <td>0 / 1 / N/A</td>
    <td>punishment, costly punishment, peer punishment, altruistic punishment, sanction, deduction</td></tr>
<tr><td><code>CONFIG_punishmentCost</code></td>
    <td>Number / N/A / N/R</td>
    <td>cost of punishment, punishment points, deduction tokens, punishment fee, spend X to punish</td></tr>
<tr><td><code>CONFIG_punishmentTech</code></td>
    <td>Number / N/A / N/R</td>
    <td>deduction ratio, points deducted per token, punishment effectiveness, each point costs X reduces by Y</td></tr>
<tr><td><code>CONFIG_showPunishmentId</code></td>
    <td>0 / 1 / N/A</td>
    <td>identified punishment, anonymous punishment, punisher identity, who punished, targeted punishment</td></tr>
<tr><td><code>CONFIG_rewardExists</code></td>
    <td>0 / 1 / N/A</td>
    <td>reward, rewarding, bonus, positive sanction, commendation, praise, applause<br><em>N/A if not mentioned anywhere</em></td></tr>
<tr><td><code>CONFIG_rewardCost</code></td>
    <td>Number / N/A / N/R</td>
    <td>cost of rewarding, reward tokens, reward fee, spend X to reward</td></tr>
<tr><td><code>CONFIG_rewardTech</code></td>
    <td>Number / N/A / N/R</td>
    <td>reward effectiveness, bonus ratio, points gained per reward token, each point costs X adds Y</td></tr>
<tr><td><code>CONFIG_showRewardId</code></td>
    <td>0 / 1 / N/A</td>
    <td>identified reward, anonymous reward, rewarder identity, who rewarded</td></tr>
</tbody></table>"""

    granularity_html = """
<ul>
<li>One <strong>condition tab</strong> per treatment arm (e.g., punishment vs. no-punishment; cohorts with different parameters).</li>
<li>Conditions within a paper share study-type fields but may differ in every CONFIG_ field.</li>
<li>Label each tab clearly (e.g., "Punishment", "Baseline", "Study 2 — Reward").</li>
</ul>"""

    workflow_html = """
<ol style="padding-left:18px;line-height:2.2">
<li>Determine <strong>Empirical</strong>, <strong>Controlled_Or_Observational</strong>, and <strong>Lab_Or_Field</strong> from the abstract and methods.</li>
<li>If <code>Empirical = 0</code> or <code>Lab_Or_Field = 1</code>, all CONFIG_ fields are N/A automatically — add a note in Misc if useful and move on.</li>
<li>Identify each distinct condition and create one tab per condition.</li>
<li>Fill CONFIG_ fields from the design/methods section. Hover over any field label in the app to see its definition.</li>
<li>Use <strong>N/R</strong> when the value should exist but is not reported; <strong>N/A</strong> when the concept does not apply to this condition.</li>
<li>Record table/figure references and edge-case notes in <strong>Misc</strong>.</li>
</ol>"""

    pitfalls_html = """
<table class="na-table">
<thead><tr><th style="width:33%">Field</th><th style="width:28%">Common mistake</th><th>Correct rule</th></tr></thead>
<tbody>
<tr><td><code>CONFIG_allOrNothing</code></td>
    <td>Coding 1 = all-or-nothing</td>
    <td><strong>1 = continuous</strong> (any specific amount); 0 = binary. The field name is counterintuitive.</td></tr>
<tr><td><code>CONFIG_chat</code></td>
    <td>Coding 0 when not mentioned</td>
    <td>Silence → N/A. Code 0 only if the paper <em>explicitly</em> prohibits communication.</td></tr>
<tr><td><code>CONFIG_showOtherSummaries</code></td>
    <td>Coding 1 when contributions are visible</td>
    <td>Contribution visibility ≠ payoff visibility. Only individual <em>payoffs/earnings</em> visible to others → 1.</td></tr>
<tr><td><code>CONFIG_punishmentExists</code> / <code>CONFIG_rewardExists</code></td>
    <td>Inferring from overall paper design</td>
    <td>Code per condition independently. If not described for this condition → N/A.</td></tr>
<tr><td><code>CONFIG_rewardExists</code></td>
    <td>Coding 0 when not mentioned</td>
    <td>Silence → N/A, not 0. Code 0 only if the paper explicitly rules out rewards.</td></tr>
<tr><td><code>CONFIG_punishmentTech</code> / <code>CONFIG_rewardTech</code></td>
    <td>Using raw magnitude</td>
    <td>Value = magnitude ÷ cost. E.g., punisher spends 1, target loses 3 → value is 3.</td></tr>
<tr><td><code>CONFIG_playerCount</code></td>
    <td>Using total participant count</td>
    <td>Count players in <em>one group</em>, not all participants across all groups.</td></tr>
<tr><td><code>CONFIG_showNRounds</code></td>
    <td>Inferring from fixed round count</td>
    <td>Code 1 only if the display is <em>explicitly stated</em>. Fixed/known rounds ≠ displayed. Silence → N/A.</td></tr>
</tbody></table>"""

    mpcr_formula_html = '<div class="formula-box">MPCR = Fund Multiplier ÷ Group Size</div>'

    # ── Build field nav + cards ─────────────────────────────────────────────
    field_nav = "\n".join(
        f'<a href="#{html_mod.escape(field_anchor(f["name"]))}" class="nav-link">'
        f'{html_mod.escape(f["name"])}</a>'
        for f in fields
    )

    field_cards = ""
    for f in fields:
        anchor = field_anchor(f["name"])
        notes_html = f'<div class="card-notes">{html_mod.escape(f["notes"])}</div>' if f["notes"] else ""
        syns = ", ".join(
            f'<span class="syn">{html_mod.escape(s.strip())}</span>'
            for s in f["synonyms"].split(",") if s.strip()
        ) if f["synonyms"] else ""
        syn_block = f'<div class="card-syns"><span class="card-label">Synonyms:</span> {syns}</div>' if syns else ""
        field_cards += f"""
<div class="field-card" id="{html_mod.escape(anchor)}">
  <div class="card-header">
    <code class="field-name">{html_mod.escape(f["name"])}</code>
    <span class="field-type">{html_mod.escape(f["type"])}</span>
  </div>
  <div class="card-def">{html_mod.escape(f["definition"])}</div>
  {syn_block}
  {notes_html}
</div>"""

    na_table_html = ""
    if na_rows:
        rows_html = "\n".join(
            f"<tr><td><strong>{html_mod.escape(r[0])}</strong></td>"
            f"<td>{html_mod.escape(r[1])}</td>"
            f"<td><em>{html_mod.escape(r[2])}</em></td></tr>"
            for r in na_rows
        )
        na_table_html = f"""
<table class="na-table">
<thead><tr><th>Value</th><th>When to use</th><th>Example</th></tr></thead>
<tbody>{rows_html}</tbody>
</table>"""

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Extraction Guide — PGG</title>
<style>
*, *::before, *::after {{ box-sizing: border-box; margin: 0; padding: 0; }}
:root {{
  --accent: #2563eb; --accent-lt: #eff6ff; --border: #e4e4e7;
  --text: #18181b; --muted: #71717a; --surface: #fff; --bg: #f4f4f5;
  --nav-w: 220px;
}}
body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
        background: var(--bg); color: var(--text); font-size: 14px; }}
#wrap {{ display: flex; min-height: 100vh; }}
#nav {{
  width: var(--nav-w); flex-shrink: 0; background: #1e3a8a; color: #fff;
  position: sticky; top: 0; height: 100vh; overflow-y: auto;
  padding: 16px 0; display: flex; flex-direction: column; gap: 2px;
}}
#nav h2 {{ font-size: 11px; font-weight: 700; text-transform: uppercase;
            letter-spacing: 1px; color: rgba(255,255,255,.5); padding: 10px 14px 4px; margin-top: 6px; }}
.nav-link {{
  display: block; padding: 5px 14px; font-size: 12px; font-weight: 500;
  color: rgba(255,255,255,.75); text-decoration: none;
  font-family: 'Menlo','Monaco',monospace;
  white-space: nowrap; overflow: hidden; text-overflow: ellipsis;
  transition: background .1s, color .1s;
}}
.nav-link:hover {{ background: rgba(255,255,255,.12); color: #fff; }}
.nav-sep {{ height: 1px; background: rgba(255,255,255,.12); margin: 8px 14px; }}
.nav-plain {{ padding: 5px 14px; font-size: 12px; color: rgba(255,255,255,.75); text-decoration: none; display: block; }}
.nav-plain:hover {{ background: rgba(255,255,255,.1); color: #fff; }}
#main {{ flex: 1; max-width: 860px; padding: 40px 48px 80px; }}
h1 {{ font-size: 26px; font-weight: 800; color: #1e3a8a; margin-bottom: 6px; }}
.subtitle {{ color: var(--muted); font-size: 14px; margin-bottom: 32px; }}
h2.sec {{ font-size: 18px; font-weight: 700; color: #1e3a8a;
           border-bottom: 2px solid var(--accent); padding-bottom: 6px; margin: 40px 0 16px; }}
h3.subsec {{ font-size: 15px; font-weight: 700; color: #1e40af; margin: 24px 0 8px; }}
p {{ font-size: 13.5px; line-height: 1.75; color: #374151; margin-bottom: 10px; }}
ul, ol {{ margin: 8px 0 12px 20px; }}
li {{ font-size: 13.5px; line-height: 1.65; color: #374151; margin-bottom: 4px; }}
code {{ background: #f1f5f9; padding: 1px 5px; border-radius: 3px;
         font-size: 12px; font-family: 'Menlo','Monaco',monospace; }}
.field-card {{
  background: var(--surface); border: 1px solid var(--border);
  border-radius: 8px; padding: 16px 18px; margin-bottom: 12px; scroll-margin-top: 20px;
}}
.field-card:target {{ border-color: var(--accent); box-shadow: 0 0 0 3px var(--accent-lt); }}
.card-header {{ display: flex; align-items: center; gap: 12px; margin-bottom: 8px; }}
.field-name {{
  font-size: 14px; font-weight: 700; font-family: 'Menlo','Monaco',monospace;
  background: var(--accent-lt); color: var(--accent); padding: 2px 8px; border-radius: 4px;
}}
.field-type {{
  font-size: 11px; font-weight: 600; background: #f3f4f6; color: #374151;
  padding: 2px 7px; border-radius: 4px; border: 1px solid var(--border);
}}
.card-def {{ font-size: 13.5px; line-height: 1.65; color: #374151; margin-bottom: 8px; }}
.card-syns {{ font-size: 12px; color: var(--muted); margin-bottom: 6px; }}
.card-label {{ font-weight: 600; color: #52525b; }}
.syn {{
  display: inline-block; background: #f3f4f6; border: 1px solid var(--border);
  border-radius: 3px; padding: 0 5px; margin: 1px 2px; font-size: 11px;
}}
.card-notes {{
  font-size: 12.5px; background: #fffbeb; border: 1px solid #fde68a;
  border-radius: 5px; padding: 8px 12px; color: #78350f; line-height: 1.6;
}}
.na-table {{ border-collapse: collapse; width: 100%; margin: 12px 0 20px; font-size: 13px; }}
.na-table th {{ background: #1e3a8a; color: white; padding: 8px 12px; text-align: left; }}
.na-table td {{ border: 1px solid var(--border); padding: 8px 12px; vertical-align: top; }}
.na-table tr:nth-child(even) td {{ background: var(--bg); }}
.formula-box {{
  background: #f0fdf4; border: 1px solid #86efac; border-radius: 6px;
  padding: 12px 16px; margin: 12px 0; font-family: 'Menlo','Monaco',monospace;
  font-size: 13px; font-weight: 700; color: #14532d;
}}
.back-btn {{
  display: inline-flex; align-items: center; gap: 6px;
  background: #1e3a8a; color: white; text-decoration: none;
  padding: 7px 14px; border-radius: 6px; font-size: 12px; font-weight: 600; margin-bottom: 28px;
}}
.back-btn:hover {{ background: #1e40af; }}
::-webkit-scrollbar {{ width: 5px; }}
::-webkit-scrollbar-thumb {{ background: #d4d4d8; border-radius: 10px; }}
</style>
</head>
<body>
<div id="wrap">

<nav id="nav">
  <h2>Navigation</h2>
  <a class="nav-plain" href="#overview">Overview</a>
  <a class="nav-plain" href="#na-nr">N/A vs N/R</a>
  <a class="nav-plain" href="#workflow">Workflow</a>
  <a class="nav-plain" href="#pitfalls">Pitfalls</a>
  <a class="nav-plain" href="#fields">Field Definitions</a>
  <div class="nav-sep"></div>
  <h2>Fields</h2>
  {field_nav}
</nav>

<main id="main">
  <a class="back-btn" href="/" target="_self">← Back to Extraction App</a>
  <h1>Field Extraction Guide</h1>
  <p class="subtitle">Design space parameters — PGG lab experiments</p>

  <h2 class="sec" id="overview">Overview</h2>
  {overview_html}
  <h3 class="subsec" id="granularity">One tab = one condition</h3>
  {granularity_html}

  <h2 class="sec" id="na-nr">N/A vs N/R</h2>
  {na_table_html}

  <h2 class="sec" id="workflow">Workflow</h2>
  {workflow_html}

  <h2 class="sec" id="pitfalls">Common Pitfalls</h2>
  {pitfalls_html}

  <h2 class="sec" id="fields">Field Definitions</h2>
  <p style="margin-bottom:16px">Hover over any field label in the extraction app to see its definition inline. Full definitions below.</p>
  {field_cards}

  <h3 class="subsec" id="CONFIG_MPCR">MPCR formula</h3>
  {mpcr_formula_html}
</main>
</div>
</body>
</html>"""


@app.route("/guide")
def guide():
    global _guide_cache
    if _guide_cache is None:
        _guide_cache = _build_guide_html()
    if _guide_cache is None:
        return "<h2>Guide not found</h2><p>Place the DOCX at the configured path.</p>", 404
    return _guide_cache


if __name__ == "__main__":
    app.run(debug=True, port=5050)
